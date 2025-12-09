#!/usr/bin/env python3
"""
Test script to validate all converters work correctly.
Creates sample test files and tests each converter.
"""

import os
import sys
import tempfile
from pathlib import Path

# Add project root to path
sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))

from converters import CSVConverter, PDFConverter, DOCXConverter, TXTConverter
from utils.file_utils import ensure_directory_exists


def create_test_csv():
    """Create a test CSV file."""
    csv_path = os.path.join(tempfile.gettempdir(), "test_input.csv")
    with open(csv_path, 'w') as f:
        f.write("Name,Age,City\n")
        f.write("Alice,30,New York\n")
        f.write("Bob,25,Los Angeles\n")
        f.write("Charlie,35,Chicago\n")
    print(f"✓ Created test CSV: {csv_path}")
    return csv_path


def create_test_txt():
    """Create a test text file."""
    txt_path = os.path.join(tempfile.gettempdir(), "test_input.txt")
    with open(txt_path, 'w') as f:
        f.write("Name Age City\n")
        f.write("Alice 30 NewYork\n")
        f.write("Bob 25 LosAngeles\n")
        f.write("Charlie 35 Chicago\n")
    print(f"✓ Created test TXT: {txt_path}")
    return txt_path


def create_test_docx():
    """Create a test DOCX file using python-docx."""
    from docx import Document
    
    docx_path = os.path.join(tempfile.gettempdir(), "test_input.docx")
    doc = Document()
    doc.add_heading("Test Document", level=1)
    doc.add_paragraph("This is the first paragraph.")
    doc.add_paragraph("This is the second paragraph.")
    doc.save(docx_path)
    print(f"✓ Created test DOCX: {docx_path}")
    return docx_path


def test_csv_converter():
    """Test CSV converter."""
    print("\n--- Testing CSV Converter ---")
    csv_path = create_test_csv()
    output_xlsx = os.path.join(tempfile.gettempdir(), "output_test.xlsx")
    
    try:
        converter = CSVConverter(csv_path)
        result = converter.convert(output_xlsx)
        if result and os.path.exists(output_xlsx):
            print(f"✓ CSV to XLSX conversion successful: {output_xlsx}")
            return True
        else:
            print("✗ CSV to XLSX conversion failed")
            return False
    except Exception as e:
        print(f"✗ CSV converter error: {e}")
        return False


def test_docx_converter():
    """Test DOCX converter."""
    print("\n--- Testing DOCX Converter ---")
    docx_path = create_test_docx()
    output_txt = os.path.join(tempfile.gettempdir(), "output_test.txt")
    
    try:
        converter = DOCXConverter(docx_path)
        result = converter.convert(output_txt)
        if result and os.path.exists(output_txt):
            print(f"✓ DOCX to TXT conversion successful: {output_txt}")
            return True
        else:
            print("✗ DOCX to TXT conversion failed")
            return False
    except Exception as e:
        print(f"✗ DOCX converter error: {e}")
        return False


def test_txt_converter():
    """Test TXT converter."""
    print("\n--- Testing TXT Converter ---")
    txt_path = create_test_txt()
    output_csv = os.path.join(tempfile.gettempdir(), "output_test_txt.csv")
    
    try:
        converter = TXTConverter(txt_path)
        result = converter.convert(output_csv)
        if result and os.path.exists(output_csv):
            print(f"✓ TXT to CSV conversion successful: {output_csv}")
            return True
        else:
            print("✗ TXT to CSV conversion failed")
            return False
    except Exception as e:
        print(f"✗ TXT converter error: {e}")
        return False


def main():
    """Run all tests."""
    print("=" * 50)
    print("File Converter Test Suite")
    print("=" * 50)
    
    results = []
    
    # Test CSV Converter
    results.append(("CSV Converter", test_csv_converter()))
    
    # Test DOCX Converter
    results.append(("DOCX Converter", test_docx_converter()))
    
    # Test TXT Converter
    results.append(("TXT Converter", test_txt_converter()))
    
    # Summary
    print("\n" + "=" * 50)
    print("Test Summary")
    print("=" * 50)
    
    passed = sum(1 for _, result in results if result)
    total = len(results)
    
    for name, result in results:
        status = "✓ PASS" if result else "✗ FAIL"
        print(f"{status}: {name}")
    
    print(f"\nTotal: {passed}/{total} tests passed")
    
    return passed == total


if __name__ == "__main__":
    success = main()
    sys.exit(0 if success else 1)
