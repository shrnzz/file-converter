# PDF Converter module - converts PDF files to other formats

# Import libraries for PDF handling
try:
    from pdf2docx import Converter
    HAS_PDF2DOCX = True
except ImportError:
    HAS_PDF2DOCX = False

# Import the base converter class 
from .base_converter import BaseConverter
import os

"""
Converter class for handling PDF file conversions.
    Currently converts PDF files to Word (.docx) format.
"""
class PDFConverter(BaseConverter):

    """
    Return the file formats that PDF files can be converted to.
    """
    def get_supported_formats(self):
        # PDF can be converted to this format
        return ['.docx']

    """ 
    Convert a PDF file to Word format (.docx).
    """
    def convert(self, output_path):
        try:
            # Store the output path for later use
            self.output_path = output_path
            
            # Check if the input file exists before attempting to convert
            if not self.validate_input():
                print(f"Error: Input file '{self.input_path}' does not exist.")
                return False
            
            # Verify that the output file has a .docx extension
            if not output_path.endswith('.docx'):
                print("Error: Output file must have a .docx extension")
                print(f"Supported formats: {', '.join(self.get_supported_formats())}")
                return False
            
            if not HAS_PDF2DOCX:
                print("Error: pdf2docx library is not installed")
                print("Install it with: pip install pdf2docx")
                return False
            
            print(f"Reading PDF file: {self.input_path}")
            print(f"This may take a moment depending on the file size...")
            
            # Use pdf2docx with error handling for compatibility issues
            print("Converting PDF to Word format...")
            
            try:
                # Try the standard conversion method
                converter = Converter(self.input_path)
                converter.convert(output_path, start=0, end=None)
                converter.close()
                
            except Exception as e:
                # If standard method fails due to compatibility, try alternative approach
                error_msg = str(e)
                if "get_area" in error_msg.lower() or "Rect" in error_msg:
                    print("Encountered PDF parsing issue, attempting alternative conversion...")
                    
                    # Try with PyMuPDF directly for text extraction
                    try:
                        import fitz  # PyMuPDF
                        from docx import Document
                        
                        # Extract text from PDF
                        pdf_text = ""
                        pdf_doc = fitz.open(self.input_path)
                        for page_num in range(len(pdf_doc)):
                            page = pdf_doc[page_num]
                            pdf_text += f"--- Page {page_num + 1} ---\n"
                            pdf_text += page.get_text() + "\n"
                        pdf_doc.close()
                        
                        # Create Word document with extracted text
                        doc = Document()
                        doc.add_heading('PDF Content', level=1)
                        doc.add_paragraph(pdf_text)
                        doc.save(output_path)
                        
                        print(f"Conversion successful (text-based)! File saved to: {output_path}")
                        return True
                        
                    except Exception as alt_e:
                        print(f"Alternative conversion failed: {str(alt_e)}")
                        raise
                else:
                    raise
            
            print(f"Conversion successful! File saved to: {output_path}")
            return True
            
        except Exception as e:
            # If any error occurs during conversion, catch and print the error message
            error_msg = str(e)
            print(f"Error during PDF conversion: {error_msg}")
            
            # Provide helpful guidance based on error type
            if "get_area" in error_msg.lower() or "Rect" in error_msg:
                print("\nNote: This is a compatibility issue with pdf2docx library.")
                print("Tried to extract text from PDF instead.")
                print("For better PDF conversion, try: pip install --upgrade pdf2docx PyMuPDF")
            
            return False
